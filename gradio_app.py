#!/usr/bin/env python3

import asyncio
import os
import tempfile
from datetime import datetime
from io import BytesIO

import gradio as gr
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from src.workflow import StartupWorkflow

load_dotenv()

class ReportGenerator:
    """Handles generation of all report formats with clean, minimal formatting."""

    def __init__(self, result):
        self.result = result
        self.startup_idea = result.startup_idea
        self.timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    def generate_all_reports(self):
        """Generate all report formats and return them as a dictionary."""
        report_text = self._generate_text_report()
        return {
            "text": report_text,
            "pdf": self._generate_pdf_report(report_text),
            "word": self._generate_word_report(),
        }

    def _generate_text_report(self):
        """Generate a clean, minimal text report."""
        sections = [
            "STARTUP IDEA ANALYSIS REPORT",
            "=" * 50,
            f"Generated: {self.timestamp}",
            f"Analysis ID: {hash(self.startup_idea.name) % 10000:04d}",
            "",
            "EXECUTIVE SUMMARY",
            "-" * 20,
            f"Startup: {self.startup_idea.name}",
            f"Description: {self.startup_idea.description or 'No description provided'}",
            "",
            self._format_market_analysis(),
            "",
            self._format_competitors(),
            "",
            self._format_viability_assessment(),
            "",
            "RECOMMENDATIONS",
            "-" * 15,
            self._format_recommendations(),
            "",
            "KEY TAKEAWAYS",
            "-" * 13,
            self._format_key_takeaways(),
        ]
        return "\n".join(str(section) for section in sections if section is not None)

    def _format_market_analysis(self):
        """Format market analysis section."""
        market = self.startup_idea.market_analysis
        if not market:
            return "MARKET ANALYSIS\n-" * 15 + "\nNo market analysis available."
        
        return "\n".join([
            "MARKET ANALYSIS",
            "-" * 15,
            f"Market Size: {market.market_size or 'Unknown'}",
            f"Growth Rate: {market.growth_rate or 'Unknown'}",
            f"Target Audience: {', '.join(market.target_audience) if market.target_audience else 'N/A'}",
            f"Market Trends: {', '.join(market.market_trends) if market.market_trends else 'N/A'}",
            f"Barriers to Entry: {', '.join(market.barriers_to_entry) if market.barriers_to_entry else 'N/A'}"
        ])

    def _format_competitors(self):
        """Format competitors section."""
        if not self.startup_idea.competitors:
            return "COMPETITOR ANALYSIS\n" + "-" * 19 + "\nNo direct competitors identified."
        
        lines = [
            "COMPETITOR ANALYSIS",
            "-" * 19,
            f"Found {len(self.startup_idea.competitors)} competitors:",
            ""
        ]
        
        for i, comp in enumerate(self.startup_idea.competitors[:5], 1):
            lines.extend([
                f"{i}. {comp.name}",
                f"   Website: {comp.website}",
                f"   Business Model: {comp.business_model or 'Unknown'}",
                f"   Key Features: {', '.join(comp.key_features[:3]) if comp.key_features else 'N/A'}",
                ""
            ])
        
        return "\n".join(lines[:-1])  # Remove last empty line

    def _format_viability_assessment(self):
        """Format viability assessment section."""
        analysis = self.startup_idea.startup_analysis
        if not analysis:
            return "VIABILITY ASSESSMENT\n" + "-" * 20 + "\nNo viability assessment available."
        
        return "\n".join([
            "VIABILITY ASSESSMENT",
            "-" * 20,
            f"Viability Score: {analysis.viability_score or 'N/A'}/10",
            f"Market Opportunity: {analysis.market_opportunity or 'N/A'}",
            f"Competitive Advantages: {', '.join(analysis.competitive_advantage) if analysis.competitive_advantage else 'N/A'}",
            f"Potential Challenges: {', '.join(analysis.potential_challenges) if analysis.potential_challenges else 'N/A'}"
        ])

    def _format_recommendations(self):
        """Format recommendations section."""
        if not self.result.final_analysis:
            return "No strategic recommendations available."
        
        # Clean up the analysis text
        analysis_text = self.result.final_analysis.replace("**", "")
        return analysis_text.strip()

    def _format_key_takeaways(self):
        """Format key takeaways section."""
        if not self.result.recommendations:
            return "No key takeaways available."
        
        return "\n".join(f"{i}. {rec}" for i, rec in enumerate(self.result.recommendations, 1))

    def _generate_pdf_report(self, report_text):
        """Generate a clean PDF report."""
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, 
                              leftMargin=72, rightMargin=72, 
                              topMargin=72, bottomMargin=72)
        
        styles = getSampleStyleSheet()
        
        # Customize styles for clean appearance
        title_style = styles['Title']
        title_style.fontSize = 18
        title_style.spaceAfter = 20
        
        heading_style = styles['Heading1']
        heading_style.fontSize = 14
        heading_style.spaceAfter = 12
        heading_style.spaceBefore = 18
        
        normal_style = styles['Normal']
        normal_style.fontSize = 11
        normal_style.spaceAfter = 6
        
        story = []
        lines = report_text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                story.append(Spacer(1, 6))
            elif line.startswith('='):
                continue  # Skip separator lines
            elif line == "STARTUP IDEA ANALYSIS REPORT":
                story.append(Paragraph(line, title_style))
            elif line in ["EXECUTIVE SUMMARY", "MARKET ANALYSIS", "COMPETITOR ANALYSIS", 
                         "VIABILITY ASSESSMENT", "RECOMMENDATIONS", "KEY TAKEAWAYS"]:
                story.append(Paragraph(line, heading_style))
            elif line.startswith('-'):
                continue  # Skip separator lines
            else:
                story.append(Paragraph(line, normal_style))
        
        doc.build(story)
        buffer.seek(0)
        return buffer

    def _generate_word_report(self):
        """Generate a clean Word document."""
        document = Document()
        
        # Set margins
        sections = document.sections
        for section in sections:
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
        
        # Title
        title = document.add_heading('STARTUP IDEA ANALYSIS REPORT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        document.add_paragraph(f'Generated: {self.timestamp}')
        document.add_paragraph(f'Analysis ID: {hash(self.startup_idea.name) % 10000:04d}')
        
        # Executive Summary
        document.add_heading('EXECUTIVE SUMMARY', level=1)
        document.add_paragraph(f'Startup: {self.startup_idea.name}')
        document.add_paragraph(f'Description: {self.startup_idea.description or "No description provided"}')
        
        # Market Analysis
        self._add_market_section_to_doc(document)
        
        # Competitor Analysis
        self._add_competitor_section_to_doc(document)
        
        # Viability Assessment
        self._add_viability_section_to_doc(document)
        
        # Recommendations
        document.add_heading('RECOMMENDATIONS', level=1)
        if self.result.final_analysis:
            document.add_paragraph(self.result.final_analysis.replace("**", ""))
        
        # Key Takeaways
        document.add_heading('KEY TAKEAWAYS', level=1)
        if self.result.recommendations:
            for i, rec in enumerate(self.result.recommendations, 1):
                document.add_paragraph(f'{i}. {rec}')
        
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer

    def _add_market_section_to_doc(self, document):
        """Add market analysis section to Word document."""
        document.add_heading('MARKET ANALYSIS', level=1)
        market = self.startup_idea.market_analysis
        
        if not market:
            document.add_paragraph('No market analysis available.')
            return
        
        document.add_paragraph(f'Market Size: {market.market_size or "Unknown"}')
        document.add_paragraph(f'Growth Rate: {market.growth_rate or "Unknown"}')
        document.add_paragraph(f'Target Audience: {", ".join(market.target_audience) if market.target_audience else "N/A"}')
        document.add_paragraph(f'Market Trends: {", ".join(market.market_trends) if market.market_trends else "N/A"}')
        document.add_paragraph(f'Barriers to Entry: {", ".join(market.barriers_to_entry) if market.barriers_to_entry else "N/A"}')

    def _add_competitor_section_to_doc(self, document):
        """Add competitor analysis section to Word document."""
        document.add_heading('COMPETITOR ANALYSIS', level=1)
        
        if not self.startup_idea.competitors:
            document.add_paragraph('No direct competitors identified.')
            return
        
        document.add_paragraph(f'Found {len(self.startup_idea.competitors)} competitors:')
        
        for i, comp in enumerate(self.startup_idea.competitors[:5], 1):
            p = document.add_paragraph()
            p.add_run(f'{i}. {comp.name}').bold = True
            document.add_paragraph(f'Website: {comp.website}')
            document.add_paragraph(f'Business Model: {comp.business_model or "Unknown"}')
            document.add_paragraph(f'Key Features: {", ".join(comp.key_features[:3]) if comp.key_features else "N/A"}')

    def _add_viability_section_to_doc(self, document):
        """Add viability assessment section to Word document."""
        document.add_heading('VIABILITY ASSESSMENT', level=1)
        analysis = self.startup_idea.startup_analysis
        
        if not analysis:
            document.add_paragraph('No viability assessment available.')
            return
        
        document.add_paragraph(f'Viability Score: {analysis.viability_score or "N/A"}/10')
        document.add_paragraph(f'Market Opportunity: {analysis.market_opportunity or "N/A"}')
        document.add_paragraph(f'Competitive Advantages: {", ".join(analysis.competitive_advantage) if analysis.competitive_advantage else "N/A"}')
        document.add_paragraph(f'Potential Challenges: {", ".join(analysis.potential_challenges) if analysis.potential_challenges else "N/A"}')

def validate_environment():
    """Validate required environment variables."""
    required = ['GOOGLE_API_KEY', 'SERP_API_KEY']
    return [var for var in required if not os.getenv(var)]

def create_html_section(title, content):
    """Create a clean HTML section with dark theme."""
    return f"""
    <div style="background-color: #2D3748; padding: 1rem; border-radius: 6px; margin: 0.5rem 0; border-left: 4px solid #63B3ED; color: #E2E8F0;">
        <h4 style="margin-top: 0; color: #FFFFFF;">{title}</h4>
        {content}
    </div>
    """

def format_viability_score(score):
    """Format viability score with simple styling."""
    if not score:
        return "<p>No viability score available.</p>"
    
    color = "#28a745" if score >= 7 else "#ffc107" if score >= 4 else "#dc3545"
    return f"""
    <div style="background-color: {color}; color: white; font-size: 1.5rem; font-weight: bold; 
                text-align: center; padding: 1rem; border-radius: 6px; margin: 1rem 0;">
        Viability Score: {score}/10
    </div>
    """

def format_results_to_html(result):
    """Convert analysis result to clean HTML."""
    startup = result.startup_idea
    
    # Main idea section
    idea_html = f"""
    <div style="background-color: #2C5282; padding: 1rem; border-radius: 6px; margin: 1rem 0; border-left: 4px solid #63B3ED; color: #EBF8FF;">
        <h3 style="margin-top: 0; color: #FFFFFF;">Startup Idea: {startup.name}</h3>
        <p>{startup.description or 'No description provided'}</p>
    </div>
    """

    # Market analysis
    market = startup.market_analysis
    market_content = "<p style='color: #E2E8F0;'>No market analysis available.</p>"
    if market:
        market_content = f"""
        <p style='color: #E2E8F0;'><strong style='color: #FFFFFF;'>Market Size:</strong> {market.market_size or 'Unknown'}</p>
        <p style='color: #E2E8F0;'><strong style='color: #FFFFFF;'>Growth Rate:</strong> {market.growth_rate or 'Unknown'}</p>
        <p style='color: #E2E8F0;'><strong style='color: #FFFFFF;'>Target Audience:</strong> {', '.join(market.target_audience) if market.target_audience else 'N/A'}</p>
        <p style='color: #E2E8F0;'><strong style='color: #FFFFFF;'>Market Trends:</strong> {', '.join(market.market_trends) if market.market_trends else 'N/A'}</p>
        """
    market_html = create_html_section("Market Analysis", market_content)
    
    # Competitors
    competitors_content = "<p style='color: #E2E8F0;'>No competitors found.</p>"
    if startup.competitors:
        competitors_list = []
        for i, comp in enumerate(startup.competitors[:5], 1):
            competitors_list.append(f"""
            <div style="border-bottom: 1px solid #4A5568; padding: 10px 0;">
                <strong style="color: #FFFFFF;">{i}. {comp.name}</strong><br>
                <small style="color: #E2E8F0;">Website: <a href="{comp.website}" target="_blank" style="color: #63B3ED;">{comp.website}</a></small><br>
                <small style="color: #E2E8F0;">Key Features: {', '.join(comp.key_features[:3]) if comp.key_features else 'N/A'}</small>
            </div>
            """)
        competitors_content = ''.join(competitors_list)
    competitors_html = create_html_section("Competitor Analysis", competitors_content)

    # Viability assessment
    analysis = startup.startup_analysis
    viability_content = "<p style='color: #E2E8F0;'>No viability assessment available.</p>"
    if analysis:
        viability_content = f"""
        {format_viability_score(analysis.viability_score)}
        <p style='color: #E2E8F0;'><strong style='color: #FFFFFF;'>Market Opportunity:</strong> {analysis.market_opportunity or 'N/A'}</p>
        <p style='color: #E2E8F0;'><strong style='color: #FFFFFF;'>Competitive Advantages:</strong> {', '.join(analysis.competitive_advantage) if analysis.competitive_advantage else 'N/A'}</p>
        <p style='color: #E2E8F0;'><strong style='color: #FFFFFF;'>Potential Challenges:</strong> {', '.join(analysis.potential_challenges) if analysis.potential_challenges else 'N/A'}</p>
        """
    viability_html = create_html_section("Viability Assessment", viability_content)

    # Recommendations
    recommendations_content = f"<p style='color: #E2E8F0;'>{result.final_analysis or 'No recommendations available.'}</p>"
    if result.recommendations:
        recommendations_content += "<ul style='color: #E2E8F0;'>" + ''.join(f'<li>{rec}</li>' for rec in result.recommendations) + "</ul>"
    recommendations_html = create_html_section("Recommendations & Key Takeaways", recommendations_content)

    return idea_html + market_html + competitors_html + viability_html + recommendations_html

async def analyze_startup_idea(startup_idea_str, progress=gr.Progress()):
    """Main analysis function."""
    if not startup_idea_str.strip():
        error_msg = "Please enter a startup idea to analyze."
        return error_msg, *[gr.update(value=None, visible=False)] * 4

    missing_vars = validate_environment()
    if missing_vars:
        error_msg = f"Missing required environment variables: {', '.join(missing_vars)}. Check your .env file."
        return error_msg, *[gr.update(value=None, visible=False)] * 4

    progress(0, desc="Starting analysis...")
    
    try:
        workflow = StartupWorkflow()
        progress(0.3, desc="Analyzing market and competitors...")
        result = await workflow.run(startup_idea_str.strip())

        if not result or not result.startup_idea:
            error_msg = "Analysis failed. Please try again."
            return error_msg, *[gr.update(value=None, visible=False)] * 4

        progress(0.7, desc="Generating reports...")
        
        results_html = format_results_to_html(result)
        report_generator = ReportGenerator(result)
        reports = report_generator.generate_all_reports()

        # Create temporary files
        files = {}
        for format_name, content in [
            ("txt", reports["text"]),
            ("pdf", reports["pdf"].getvalue()),
            ("docx", reports["word"].getvalue())
        ]:
            with tempfile.NamedTemporaryFile(delete=False, 
                                           suffix=f".{format_name}", 
                                           mode="wb" if format_name != "txt" else "w",
                                           encoding="utf-8" if format_name == "txt" else None) as tmp:
                tmp.write(content)
                files[format_name] = tmp.name
        
        progress(1, desc="Complete!")

        return (
            results_html,
            gr.update(value=reports["text"], visible=True),
            gr.update(value=files["txt"], visible=True),
            gr.update(value=files["pdf"], visible=True),
            gr.update(value=files["docx"], visible=True),
        )

    except Exception as e:
        error_msg = f"An error occurred: {str(e)}"
        return error_msg, *[gr.update(value=None, visible=False)] * 4

def create_interface():
    """Create the Gradio interface."""
    with gr.Blocks(theme=gr.themes.Soft(), title="Startup Idea Analyzer") as demo:
        gr.Markdown("# Startup Idea Analyzer")
        gr.Markdown("AI-powered market research and competitor analysis for your startup ideas.")
        
        with gr.Row():
            with gr.Column(scale=2):
                startup_idea_input = gr.Textbox(
                    label="Describe your startup idea:",
                    placeholder="e.g., A subscription box for eco-friendly pet toys",
                    lines=4,
                )
                analyze_btn = gr.Button("Analyze Idea", variant="primary", size="lg")
            
            with gr.Column(scale=1):
                gr.Markdown("### Configuration Status")
                missing_vars = validate_environment()
                status = "Environment configured" if not missing_vars else f"Missing: {', '.join(missing_vars)}"
                gr.Markdown(f"**{status}**")
                
                gr.Markdown("### About")
                gr.Markdown("Comprehensive startup analysis covering market size, competition, and viability assessment.")
        
        results_html = gr.HTML(label="Analysis Results")
        
        with gr.Row():
            with gr.Column(scale=2):
                gr.Markdown("### Report Preview")
                report_preview = gr.Textbox(
                    label="Text Report",
                    lines=20,
                    max_lines=30,
                    interactive=False,
                    visible=False
                )
            
            with gr.Column(scale=1):
                gr.Markdown("### Download Reports")
                txt_download = gr.File(label="Download Text Report", visible=False)
                pdf_download = gr.File(label="Download PDF Report", visible=False)
                word_download = gr.File(label="Download Word Document", visible=False)

        analyze_btn.click(
            fn=analyze_startup_idea,
            inputs=[startup_idea_input],
            outputs=[results_html, report_preview, txt_download, pdf_download, word_download],
        )
    
    return demo

if __name__ == "__main__":
    app = create_interface()
    app.launch()